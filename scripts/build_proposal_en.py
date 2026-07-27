"""Build the polished English DOCX proposal from proposal-content-en.md."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/proposal-content-en.md"
OUTPUT = ROOT / "Edcosys_CCTV_Retail_Loss_Prevention_Proposal_EN.docx"

FONT_LATIN = "Calibri"
FONT_CJK = "Microsoft JhengHei"
FONT_MONO = "Consolas"

BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
TEAL = "138A78"
AMBER = "B66B00"
INK = "17242C"
GRAY = "5D6B75"
MUTED = "7A8790"
LIGHT = "F2F4F7"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E7F5F2"
LIGHT_AMBER = "FFF4DE"
BORDER = "CAD3DA"
WHITE = "FFFFFF"

DIAGRAMS = [
    (ROOT / "docs/diagrams/en/01-user-workflow-en.png", "Figure 1 | Store associate event workflow"),
    (ROOT / "docs/diagrams/en/02-sidecar-architecture-en.png", "Figure 2 | Sidecar integration with the existing CCTV system"),
    (ROOT / "docs/diagrams/en/03-model-network-en.png", "Figure 3 | YOLO26 perception and event network"),
    (ROOT / "docs/diagrams/en/05-data-loop-en.png", "Figure 4 | Data and model feedback loop"),
    (ROOT / "docs/diagrams/en/04-system-architecture-en.png", "Figure 5 | Store, edge, and control-plane architecture"),
]


def create_decimal_numbering(document: Document) -> int:
    """Create a single-level decimal list that restarts at 1."""
    numbering = document.part.numbering_part.element

    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    abstract_id = max(abstract_ids, default=-1) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{abstract_id + 1:08X}"[-8:])
    abstract.append(nsid)

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)

    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal")
    level.append(number_format)

    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)

    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "right")
    level.append(level_justification)

    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "279")
    paragraph_properties.append(indent)
    level.append(paragraph_properties)

    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_CJK)
    run_properties.append(fonts)
    level.append(run_properties)

    abstract.append(level)
    first_number = numbering.find(qn("w:num"))
    if first_number is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_number), abstract)

    number_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    number_id = max(number_ids, default=0) + 1

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(number_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return number_id


def apply_numbering(paragraph, number_id: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    existing = paragraph_properties.find(qn("w:numPr"))
    if existing is not None:
        paragraph_properties.remove(existing)
    number_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number_properties.append(level)
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(number_id))
    number_properties.append(number)
    paragraph_properties.append(number_properties)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 0) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_indent = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_indent)
    tbl_indent.set(qn("w:w"), str(indent_dxa))
    tbl_indent.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_run_font(
    run,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    name: str = FONT_LATIN,
    cjk: str = FONT_CJK,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), cjk)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_border(paragraph, color: str = BLUE, size: int = 12, space: int = 4) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_hyperlink(paragraph, text: str, url: str, size: float = 10.5) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_LATIN)
    r_fonts.set(qn("w:hAnsi"), FONT_LATIN)
    r_fonts.set(qn("w:eastAsia"), FONT_CJK)
    size_node = OxmlElement("w:sz")
    size_node.set(qn("w:val"), str(int(size * 2)))
    r_pr.extend([r_fonts, color, underline, size_node])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(
    r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\)|\*[^*]+\*)"
)


def add_inline(paragraph, text: str, size: float = 11, color: str = INK) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=max(8.2, size - 0.3), color=BLUE_DARK, name=FONT_MONO, cjk=FONT_CJK)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "EEF2F5")
            run._element.get_or_add_rPr().append(shading)
        elif token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                add_hyperlink(paragraph, link_match.group(1), link_match.group(2), size=size)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, italic=True, color=color)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_after = Pt(6)
    pf.line_spacing = 1.10
    pf.widow_control = True

    heading_specs = {
        "Title": (26, INK, 0, 6),
        "Subtitle": (14, GRAY, 0, 10),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, BLUE_DARK, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    caption = styles["Caption"]
    caption.font.name = FONT_LATIN
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    caption.font.size = Pt(8.5)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(GRAY)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(9)


def configure_section(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_header_footer(section) -> None:
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_width(table, 9360, 0)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = left.add_run("EDCOSYS  |  CCTV RETAIL LOSS PREVENTION")
    set_run_font(run, size=8.2, bold=True, color=BLUE)
    run = right.add_run("PROPOSAL  v1.0  |  27 JULY 2026")
    set_run_font(run, size=8.2, color=GRAY)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 0, 0, 0, 0)
    rule = header.add_paragraph()
    rule.paragraph_format.space_after = Pt(0)
    set_paragraph_border(rule, color=BORDER, size=4, space=1)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_width(table, 9360, 0)
    left = table.cell(0, 0).paragraphs[0]
    right = table.cell(0, 1).paragraphs[0]
    run = left.add_run("Edcosys | Technical and Product Proposal")
    set_run_font(run, size=8, color=MUTED)
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = right.add_run("Page ")
    set_run_font(run, size=8, color=MUTED)
    page_run = right.add_run()
    set_run_font(page_run, size=8, color=MUTED)
    add_field(page_run, "PAGE")
    for cell in table.rows[0].cells:
        set_cell_margins(cell, 0, 0, 0, 0)


def add_spacer(document: Document, points: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(points)


def add_cover(document: Document) -> None:
    add_spacer(document, 56)

    author = document.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("Edcosys")
    set_run_font(run, size=12, bold=True, color=GRAY)
    author.paragraph_format.space_after = Pt(8)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(title, "CCTV Retail Loss Prevention", size=24, color="000000")
    title.paragraph_format.space_after = Pt(4)

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(
        subtitle,
        "YOLO26 Feasibility, Dataset Strategy, and Pilot Proposal",
        size=14,
        color=GRAY,
    )
    subtitle.paragraph_format.space_after = Pt(8)

    descriptor = document.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_inline(
        descriptor,
        "Technical and Product Proposal | Existing CCTV Sidecar Deployment",
        size=10.5,
        color=GRAY,
    )
    for run in descriptor.runs:
        run.bold = True
    descriptor.paragraph_format.space_after = Pt(26)

    rule = document.add_paragraph()
    set_paragraph_border(rule, color=BLUE, size=18, space=4)
    rule.paragraph_format.space_after = Pt(13)

    metadata_left = [
        ("Engagement", "Measured retail loss-prevention pilot"),
        ("Scope", "One store / 2–4 cameras"),
        ("Integration", "Existing NVR + AI edge sidecar"),
        ("Operating mode", "Shadow pilot + staff review"),
    ]
    metadata_right = [
        ("Version", "v1.0 English Proposal"),
        ("Date", "27 July 2026"),
        ("Model stack", "YOLO26 + tracking + temporal model"),
        ("Decision", "Proceed with a controlled pilot"),
    ]
    table = document.add_table(rows=4, cols=4)
    set_table_width(table, 9360, 0)
    table.autofit = False
    widths = [1050, 3500, 1050, 3760]
    for row_index in range(4):
        values = (
            metadata_left[row_index][0],
            metadata_left[row_index][1],
            metadata_right[row_index][0],
            metadata_right[row_index][1],
        )
        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.width = Inches(widths[col_index] / 1440)
            set_cell_margins(cell, 70, 100, 70, 100)
            if col_index in (0, 2):
                set_cell_shading(cell, LIGHT)
            paragraph = cell.paragraphs[0]
            add_inline(
                paragraph,
                value,
                size=8.7 if col_index in (0, 2) else 9.2,
                color=GRAY if col_index in (0, 2) else INK,
            )
            if col_index in (0, 2):
                for run in paragraph.runs:
                    run.bold = True

    add_spacer(document, 4)
    callout = document.add_table(rows=1, cols=1)
    set_table_width(callout, 9360, 0)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell, 140, 170, 140, 170)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("Operating principle  ")
    set_run_font(run, size=10.2, bold=True, color=BLUE_DARK)
    add_inline(
        paragraph,
        "The system creates reviewable alerts from observable actions. Store staff use the evidence clip, live context, and store procedures to decide the response.",
        size=10.2,
        color=INK,
    )
    paragraph.paragraph_format.line_spacing = 1.10

    add_spacer(document, 0)
    metric = document.add_table(rows=1, cols=3)
    set_table_width(metric, 9360, 0)
    set_row_cant_split(metric.rows[0])
    items = [
        ("PROTOTYPE", "Interactive web UI"),
        ("LOCAL VALIDATION", "RTX 4060 real-video run"),
        ("FIRST DEPLOYMENT", "10–14 week pilot"),
    ]
    for index, (label, value) in enumerate(items):
        cell = metric.cell(0, index)
        set_cell_shading(cell, LIGHT_BLUE if index != 1 else LIGHT_TEAL)
        set_cell_margins(cell, 120, 120, 120, 120)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label + "\n")
        set_run_font(run, size=8.3, bold=True, color=BLUE if index != 1 else TEAL)
        run = paragraph.add_run(value)
        set_run_font(run, size=10, bold=True, color=INK)

    document.add_page_break()


def add_contents(document: Document, lines: list[str]) -> None:
    heading = document.add_paragraph(style="Heading 1")
    add_inline(heading, "Contents", size=16, color=BLUE)
    intro = document.add_paragraph()
    add_inline(
        intro,
        "The proposal moves from the recommended decision to user workflow, model and data design, architecture, prototype evidence, delivery, and acceptance.",
        size=10.5,
        color=GRAY,
    )
    headings = []
    for line in lines:
        if line.startswith("## ") and re.match(r"## \d+\.", line):
            headings.append(line[3:].strip())
    split = (len(headings) + 1) // 2
    table = document.add_table(rows=max(split, len(headings) - split), cols=2)
    set_table_width(table, 9360, 0)
    for col, items in enumerate((headings[:split], headings[split:])):
        for row, item in enumerate(items):
            cell = table.cell(row, col)
            set_cell_margins(cell, 70, 120, 70, 120)
            if row % 2 == 0:
                set_cell_shading(cell, LIGHT)
            paragraph = cell.paragraphs[0]
            add_inline(paragraph, item, size=9.2, color=INK)
    add_spacer(document, 14)
    callout = document.add_table(rows=1, cols=1)
    set_table_width(callout, 9360, 0)
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT_TEAL)
    set_cell_margins(cell, 120, 150, 120, 150)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("Recommendation  ")
    set_run_font(run, size=10.3, bold=True, color=TEAL)
    add_inline(
        paragraph,
        "Approve a one-store, 2–4 camera pilot. YOLO26 supplies the visual perception layer; tracking, temporal modeling, scene rules, store data, and staff review complete the event workflow.",
        size=10.3,
        color=INK,
    )
    document.add_page_break()


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator_row(cells: list[str]) -> bool:
    return all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def table_column_widths(rows: list[list[str]], total_dxa: int = 9360) -> list[int]:
    column_count = max(len(row) for row in rows)
    maxima = []
    for column in range(column_count):
        lengths = [len(re.sub(r"[`*_\[\]()]", "", row[column])) if column < len(row) else 0 for row in rows]
        maxima.append(max(3, min(max(lengths, default=3), 70)))
    if column_count >= 5:
        minimum, maximum = 0.72, 1.9
    elif column_count == 4:
        minimum, maximum = 0.65, 2.5
    else:
        minimum, maximum = 0.55, 4.5
    weights = [max(minimum, min(maximum, value ** 0.55)) for value in maxima]
    total_weight = sum(weights)
    widths = [int(total_dxa * weight / total_weight) for weight in weights]
    widths[-1] += total_dxa - sum(widths)
    return widths


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    normalized = [row + [""] * (column_count - len(row)) for row in rows]
    table = document.add_table(rows=len(normalized), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, 9360, 120)
    widths = table_column_widths(normalized)
    font_size = 8.0 if column_count >= 5 else 8.6 if column_count == 4 else 9.1 if column_count == 3 else 9.3
    for row_index, row in enumerate(normalized):
        for column_index, text in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if row_index == 0 else WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.width = Inches(widths[column_index] / 1440)
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, LIGHT)
            elif row_index % 2 == 0:
                set_cell_shading(cell, "FAFBFC")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.12
            add_inline(paragraph, text, size=font_size, color=INK if row_index else BLUE_DARK)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
        if row_index == 0:
            set_repeat_table_header(table.rows[row_index])
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_callout(document: Document, text: str) -> None:
    tone_fill = LIGHT_AMBER if "warning" in text.lower() or "risk" in text.lower() else LIGHT_BLUE
    tone_color = AMBER if tone_fill == LIGHT_AMBER else BLUE_DARK
    table = document.add_table(rows=1, cols=1)
    set_table_width(table, 9360, 0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, tone_fill)
    set_cell_margins(cell, 110, 150, 110, 150)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.28
    add_inline(paragraph, text, size=9.8, color=tone_color)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_image_with_caption(document: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        paragraph = document.add_paragraph()
        add_inline(paragraph, f"[Missing image: {image_path}]", size=9, color=AMBER)
        return
    with Image.open(image_path) as image:
        width_px, height_px = image.size
    width_inches = 6.42
    height_inches = width_inches * height_px / width_px
    if height_inches > 7.2:
        width_inches *= 7.2 / height_inches
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    caption_paragraph = document.add_paragraph(style="Caption")
    add_inline(caption_paragraph, caption, size=8.5, color=GRAY)


def markdown_image(line: str) -> tuple[str, Path] | None:
    match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
    if not match:
        return None
    caption = match.group(1)
    target = (SOURCE.parent / match.group(2)).resolve()
    return caption, target


def add_body(document: Document, lines: list[str]) -> None:
    start = next(index for index, line in enumerate(lines) if line.startswith("## 1."))
    lines = lines[start:]
    diagram_index = 0
    image_index = 0
    i = 0
    major_page_breaks = {4, 9, 13}
    current_order_number_id: int | None = None
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue
        if stripped == "---":
            current_order_number_id = None
            i += 1
            continue

        if stripped.startswith("```mermaid"):
            current_order_number_id = None
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            i += 1
            if diagram_index < len(DIAGRAMS):
                path, caption = DIAGRAMS[diagram_index]
                add_image_with_caption(document, path, caption)
                diagram_index += 1
            continue

        if stripped.startswith("```"):
            current_order_number_id = None
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            if i < len(lines):
                i += 1
            code_table = document.add_table(rows=1, cols=1)
            set_table_width(code_table, 9360, 120)
            code_cell = code_table.cell(0, 0)
            set_cell_shading(code_cell, LIGHT)
            set_cell_margins(code_cell, 100, 140, 100, 140)
            code_paragraph = code_cell.paragraphs[0]
            code_paragraph.paragraph_format.space_after = Pt(0)
            code_paragraph.paragraph_format.line_spacing = 1.10
            code_run = code_paragraph.add_run("\n".join(code_lines))
            set_run_font(code_run, size=8.7, color=INK, name=FONT_MONO, cjk=FONT_MONO)
            continue

        image_info = markdown_image(stripped)
        if image_info:
            current_order_number_id = None
            caption, path = image_info
            image_index += 1
            add_image_with_caption(document, path, caption)
            i += 1
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            current_order_number_id = None
            level = len(heading_match.group(1)) - 1
            text = heading_match.group(2)
            section_match = re.match(r"(\d+)\.", text)
            if level == 1 and section_match and int(section_match.group(1)) in major_page_breaks:
                document.add_page_break()
            paragraph = document.add_paragraph(style=f"Heading {min(level, 3)}")
            add_inline(
                paragraph,
                text,
                size=16 if level == 1 else 13 if level == 2 else 12,
                color=BLUE if level < 3 else BLUE_DARK,
            )
            i += 1
            continue

        if stripped.startswith(">"):
            current_order_number_id = None
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip()[1:].strip())
                i += 1
            add_callout(document, " ".join(quote_lines))
            continue

        if stripped.startswith("|"):
            current_order_number_id = None
            raw_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                raw_rows.append(parse_table_row(lines[i]))
                i += 1
            rows = [row for row in raw_rows if not is_separator_row(row)]
            add_markdown_table(document, rows)
            continue

        list_match = re.match(r"^(\s*)([-*]|\d+\.)\s+(.+)$", line)
        if list_match:
            ordered = list_match.group(2).endswith(".") and list_match.group(2)[0].isdigit()
            if ordered:
                if current_order_number_id is None:
                    current_order_number_id = create_decimal_numbering(document)
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.5)
                paragraph.paragraph_format.first_line_indent = Inches(-0.25)
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.paragraph_format.line_spacing = 1.167
                apply_numbering(paragraph, current_order_number_id)
            else:
                current_order_number_id = None
                paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.keep_together = True
            add_inline(paragraph, list_match.group(3), size=11, color=INK)
            i += 1
            continue

        if not line.startswith((" ", "\t")):
            current_order_number_id = None
        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            candidate = lines[i].strip()
            if (
                not candidate
                or candidate == "---"
                or candidate.startswith("#")
                or candidate.startswith("|")
                or candidate.startswith(">")
                or candidate.startswith("```")
                or candidate.startswith("![")
                or re.match(r"^([-*]|\d+\.)\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            i += 1
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_together = False
        add_inline(paragraph, " ".join(paragraph_lines), size=11, color=INK)


def main() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    document = Document()
    configure_styles(document)
    configure_section(document.sections[0])
    add_header_footer(document.sections[0])
    core = document.core_properties
    core.title = "CCTV Retail Loss Prevention - YOLO26 Feasibility and Pilot Proposal"
    core.subject = "User needs, dataset strategy, model architecture, system design, prototype evidence, and delivery plan"
    core.author = "Edcosys"
    core.keywords = "YOLO26, retail loss prevention, CCTV, computer vision, proposal"
    core.comments = "Generated and verified by Edcosys."

    add_cover(document)
    add_contents(document, lines)
    add_body(document, lines)

    settings = document.settings
    settings.update_fields_on_open = True
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
